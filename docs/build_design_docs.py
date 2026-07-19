from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent
BLUE = '1F4D78'; LIGHT = 'E8EEF5'; GREY = 'F2F4F7'; INK = '0B2545'

def font(size, bold=False):
    try: return ImageFont.truetype('arial.ttf', size)
    except: return ImageFont.load_default()

def agent_diagram(path):
    img = Image.new('RGB', (1600, 780), 'white'); d = ImageDraw.Draw(img)
    title = font(38, True); label = font(23, True); body = font(18)
    d.text((50, 25), 'Saudagar AI - Agent Architecture and Event Flow', fill='#0B2545', font=title)
    boxes = [
        (55,150,310,300,'Flutter Client','Voice/text capture\nRealtime dashboard'),
        (390,150,650,300,'FastAPI Gateway','Validation, routing\nBackground tasks'),
        (730,65,1010,210,'Demand Capture Agent','Extract, verify, clean\nCanonicalise product'),
        (730,290,1010,435,'Demand Intelligence Agent','Aggregate 30-day events\nScore demand'),
        (730,515,1010,660,'Business Intelligence Agent','Weather, trends, festivals\nEvery 6 hours'),
        (1090,250,1370,410,'Procurement Agent','Fuse demand + context\nRecommend action'),
        (1395,170,1560,500,'Firestore','Products\nEvents\nSummaries\nInsights\nRecommendations'),
    ]
    for x1,y1,x2,y2,h,b in boxes:
        d.rounded_rectangle((x1,y1,x2,y2), 18, fill='#F4F6F9', outline='#1F4D78', width=3)
        d.text((x1+15,y1+18), h, fill='#1F4D78', font=label)
        d.multiline_text((x1+15,y1+62), b, fill='#222222', font=body, spacing=5)
    arrows=[(310,225,390,225),(650,190,730,140),(650,230,730,350),(1010,135,1395,225),(1010,360,1090,330),(1010,585,1090,370),(1370,330,1395,330),(1395,470,310,270)]
    for x1,y1,x2,y2 in arrows:
        d.line((x1,y1,x2,y2), fill='#2E74B5', width=4)
        d.polygon([(x2,y2),(x2-13,y2-7),(x2-13,y2+7)], fill='#2E74B5')
    img.save(path)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)
def set_cell_width(cell, width):
    tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn('w:tcW'))
    if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'), 'dxa')
def fix_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.autofit = False
    tblPr=table._tbl.tblPr; tblW=tblPr.first_child_found_in('w:tblW'); tblW.set(qn('w:w'),'9360'); tblW.set(qn('w:type'),'dxa')
    ind=OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); tblPr.append(ind)
    grid=table._tbl.tblGrid
    for col,w in zip(grid.gridCol_lst,widths): col.set(qn('w:w'),str(w))
    for row in table.rows:
        for c,w in zip(row.cells,widths):
            set_cell_width(c,w); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in c.paragraphs: p.paragraph_format.space_after=Pt(3)

def set_repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)
def page_number(p):
    r=p.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); r._r.append(fld)

def setup(doc, label):
    s=doc.sections[0]; s.top_margin=Inches(1); s.bottom_margin=Inches(1); s.left_margin=Inches(1); s.right_margin=Inches(1); s.header_distance=Inches(.492); s.footer_distance=Inches(.492)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string('222222'); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
    for name,size,col,before,after in [('Heading 1',16,'2E74B5',16,8),('Heading 2',13,'2E74B5',12,6),('Heading 3',12,'1F4D78',8,4)]:
        st=styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(col); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    header=s.header.paragraphs[0]; header.text=f'SAUDAGAR AI | {label.upper()}'; header.style='Normal'; header.alignment=WD_ALIGN_PARAGRAPH.LEFT; header.runs[0].font.color.rgb=RGBColor.from_string(BLUE); header.runs[0].font.bold=True; header.runs[0].font.size=Pt(9)
    footer=s.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; footer.add_run('Confidential Prototype Design | Page '); page_number(footer)

def title(doc, name, subtitle):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_before=Pt(34); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(name); r.bold=True; r.font.name='Calibri'; r.font.size=Pt(27); r.font.color.rgb=RGBColor.from_string(INK)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(20); r=p.add_run(subtitle); r.italic=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string('555555')
    t=doc.add_table(rows=3, cols=2); fix_table(t,[2700,6660]);
    for row,(a,b) in zip(t.rows,[('Document status','Submission-ready'),('Version','1.0 | 19 July 2026'),('Prepared for','Saudagar AI prototype review')]):
        row.cells[0].text=a; row.cells[1].text=b; shade(row.cells[0],GREY); row.cells[0].paragraphs[0].runs[0].font.bold=True
    doc.add_paragraph('Purpose statement', style='Heading 2')

def add_bullets(doc, items):
    for item in items: doc.add_paragraph(item, style='List Bullet')
def matrix(doc, headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; fix_table(t,widths); set_repeat_header(t.rows[0])
    for c,v in zip(t.rows[0].cells,headers): c.text=v; shade(c,LIGHT); c.paragraphs[0].runs[0].font.bold=True
    for values in rows:
        cells=t.add_row().cells
        for c,v in zip(cells,values): c.text=v
    return t

def hld():
    d=Document(); setup(d,'High-Level Design'); title(d,'High-Level Design (HLD)','Saudagar AI: Voice-led demand and procurement intelligence for kirana stores')
    d.add_paragraph('Saudagar AI converts informal customer interactions into a structured demand signal and context-aware procurement guidance. This HLD defines the solution boundary, target architecture, agent responsibilities, data flows, security posture, and future evolution. It is intentionally technology-aware but remains focused on system decisions rather than code-level mechanics.')
    d.add_paragraph('1. Business Objective',style='Heading 1')
    d.add_paragraph('Enable kirana-store operators to recognise unmet demand, reduce preventable stock-outs, and make more timely replenishment decisions without asking them to maintain complex manual records.')
    matrix(d,['Objective','Outcome indicator','Design response'],[
        ('Capture implicit demand','Customer requests become auditable events','Voice/text intake and structured extraction'),
        ('Improve availability','Recurring stock-outs are surfaced','Per-shop demand aggregation and prioritisation'),
        ('Improve order timing','Seasonal/contextual drivers influence replenishment','Business-intelligence enrichment and agent reasoning'),
        ('Keep the workflow simple','Owner can review actions in a mobile UI','Flutter dashboard with real-time Firestore reads'),
    ],[2100,2900,4360])
    d.add_paragraph('2. Scope and Principles',style='Heading 1')
    add_bullets(d,['In scope: Hindi/Hinglish or text demand capture, product canonicalisation, demand analytics, contextual recommendations, catalogue management, and feedback capture.','Out of scope for the current prototype: payment, supplier ordering, inventory-system write-back, identity/onboarding, and autonomous procurement execution.','The backend is the trusted write path. The client may consume approved Firestore projections but does not execute state transitions directly.','The agent system is assistive: recommendations are presented for shopkeeper review, not automatically sent to suppliers.'])
    d.add_paragraph('3. Logical Architecture',style='Heading 1')
    chart=ROOT/'agent_architecture.png'; agent_diagram(chart); d.add_picture(str(chart),width=Inches(6.5)); d.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph('Figure 1. Logical agent architecture. The core pipeline responds to a demand event; business context is refreshed independently and consumed by procurement reasoning.',style='Caption')
    d.add_paragraph('4. Agent Architecture',style='Heading 1')
    matrix(d,['Agent','Trigger','Responsibility','Primary output'],[
        ('Demand Capture','POST /capture-demand or /upload-audio','Extract, verify, clean and canonicalise one or more demand events','demand_events'),
        ('Demand Intelligence','Background task after capture','Aggregate recent per-shop events into frequencies, stock-out counts and demand scores','demand_summary/{shop_id}'),
        ('Business Intelligence','Startup then 6-hour scheduler','Refresh weather, trends and festival context','business_insights/latest'),
        ('Procurement','After summary update','Combine demand summary and business context into ranked guidance','recommendations/{shop_id}'),
    ],[1700,1900,3600,2160])
    d.add_paragraph('5. End-to-End Flow',style='Heading 1')
    for n,text in enumerate(['The shopkeeper records audio or provides a transcript in the Flutter client. Audio is transcribed server-side before the demand workflow starts.','FastAPI validates the request and invokes the Demand Capture pipeline. The pipeline uses model-based extraction and verification, cleans product phrases, and reconciles them with the live catalogue.','The canonical demand event is persisted. A background task derives recent demand metrics for the shop.','The Procurement Agent combines the summary with the latest seasonal and market context to generate a prioritised recommendation set.','Firestore’s real-time capability makes summaries, insights and recommendations available to the client dashboard.'],1): d.add_paragraph(text,style='List Number')
    d.add_paragraph('6. Data Architecture',style='Heading 1')
    d.add_paragraph('Firestore is the operational datastore. Events are immutable-in-principle capture records; summaries and recommendations are replaceable projections keyed by shop. The products catalogue supplies canonical names and aliases. Feedback forms a separately auditable loop for evaluating recommendation usefulness.')
    matrix(d,['Data domain','Key','Retention / mutability posture'],[
        ('Product catalogue','Product document','Managed reference data; aliases evolve with curation'),('Demand events','Generated document ID','Append-only audit signal; timestamped by backend'),('Demand summary','shop_id','Recomputed projection from recent events'),('Business insights','latest','Time-sensitive singleton refreshed on schedule'),('Recommendations','shop_id','Latest advisory projection; human review expected'),('Feedback','Generated document ID','Append-only evaluation record')],[2100,2200,5060])
    d.add_paragraph('7. Security and Trust Boundaries',style='Heading 1')
    add_bullets(d,['Service credentials and model/API keys are held only by the FastAPI runtime; they are never distributed with the Flutter client.','The client is designed for read-only realtime consumption of approved Firestore collections. Backend services own writes and validation.','Input is validated through Pydantic contracts; model outputs are verified and canonicalised before persistence.','Production deployment must replace prototype CORS wildcard settings with an explicit frontend allow-list and apply least-privilege Firestore rules.'])
    d.add_paragraph('8. Reliability, Observability and Risks',style='Heading 1')
    matrix(d,['Risk','Current mitigation','Planned hardening'],[
        ('Model uncertainty','Independent verification and transcript-grounding checks','Quality metrics, fallback review queue and prompt regression suite'),('External API failure','Error handling and local mock Firestore','Timeouts, retries, circuit breaking and durable queues'),('Duplicate/background work','Simple prototype task triggering','Idempotency keys and event-driven workers'),('Sensitive data','Server-side secret isolation','Data minimisation, retention policy and consent controls')],[2200,3500,3660])
    d.add_paragraph('9. Future Scope',style='Heading 1')
    add_bullets(d,['Integrate supplier catalogues, order drafts and inventory-system reconciliation while preserving explicit shopkeeper approval.','Introduce multilingual capture quality measurement, per-store catalogues, and active-learning alias confirmation.','Move scheduled and post-event work to durable cloud jobs/queues for horizontal scaling and retry control.','Add role-based access, tenancy isolation, audit dashboards, model evaluation, and operational SLIs/SLOs.','Use feedback to evaluate recommendation acceptance and calibrate ranking policies without allowing uncontrolled autonomous ordering.'])
    d.add_paragraph('10. Acceptance Summary',style='Heading 1'); d.add_paragraph('The HLD is satisfied when a reviewer can trace the path from customer request to a persisted demand event, derived shop-level analytics, contextual recommendation, and visible client result; understand each agent’s authority; and identify the controls required for production hardening.')
    d.save(ROOT/'HLD_Saudagar_AI.docx')

def lld():
    d=Document(); setup(d,'Low-Level Design'); title(d,'Low-Level Design (LLD)','Saudagar AI: implementation blueprint for the backend, agents, data, APIs and client integration')
    d.add_paragraph('This LLD describes the implementation represented by the repository as of version 1.0. It separates existing prototype behaviour from recommended production hardening so that implementation teams and reviewers can distinguish fact from target state.')
    d.add_paragraph('1. Implementation Topology',style='Heading 1')
    matrix(d,['Layer','Implementation','Responsibility'],[
        ('Client','Flutter / Dart','Audio capture, API invocation, localisation, dashboard streams and feedback'),('API','FastAPI + Uvicorn','HTTP contracts, CORS, lifecycle scheduling and exception mapping'),('Agents','Python service objects','Demand capture, demand aggregation, external-context refresh and procurement advice'),('Persistence','Firebase Admin SDK / Firestore','Catalogue, events, projections, recommendations and feedback'),('AI / matching','Gemini service + RapidFuzz','Structured interpretation, verification, naming cleanup and catalogue reconciliation')],[1700,3200,4460])
    d.add_paragraph('2. Backend Startup and Dependency Lifecycle',style='Heading 1')
    d.add_paragraph('`backend/main.py` instantiates the FastAPI application, configures CORS, and creates an APScheduler instance. On startup, it attempts to seed the product catalogue, warms the matching cache, runs the Business Intelligence Agent once, then schedules business intelligence every six hours and cache refresh every hour. On shutdown, the scheduler is stopped.')
    d.add_paragraph('Implementation note: FastAPI lifecycle decorators and in-process APScheduler are appropriate for the prototype. A multi-instance deployment must ensure only one scheduler leader runs, or replace the scheduler with a managed cron/job service.')
    d.add_paragraph('3. API Contracts',style='Heading 1')
    matrix(d,['Endpoint','Input','Success response / side effect'],[
        ('POST /capture-demand','shop_id, transcript','CaptureDemandResponse; stores event and schedules analytics'),('POST /upload-audio','multipart shop_id + audio file','Transcript plus CaptureDemandResponse; schedules analytics'),('POST /confirm-product-alias','shop_id, canonical_name, new_alias','Confirmation result; current service implementation requires completion'),('GET /recommendations','shop_id query','Latest shop recommendation projection'),('GET /demand-summary','shop_id query','Latest demand metric projection'),('GET /business-insights','None','Latest global business-context projection'),('GET/POST /products','None / ProductModel','Catalogue read or create; refreshes matching cache'),('POST /feedback','FeedbackRequest','Persists audit feedback'),('POST /run-bi','None','Forces business-context refresh')],[2200,3100,4060])
    d.add_paragraph('4. Demand Capture Agent - Detailed Pipeline',style='Heading 1')
    d.add_paragraph('`DemandCapturePipeline.run(shop_id, transcript)` produces zero or more persisted demand events. It is called by both text and audio entry points; the audio endpoint first calls `sarvam_service.transcribe_audio` to obtain a transcript.')
    for t in ['Stage 1 - Conversation analysis: Gemini extracts candidate demand events from the transcript. A missing/invalid `events` payload terminates the pipeline without persistence.','Stage 2 - Independent verification: each candidate is re-evaluated against the original transcript. The implementation checks verified status, confidence threshold, product presence and evidence grounding. Requests with unknown availability use a lower 0.65 threshold; other events use 0.82.','Stage 3 - Product cleaning: Gemini normalises the requested phrase. A token-subset or contiguous-substring post-check rejects unsupported transformations and falls back to the raw phrase.','Stage 4/6 - Catalogue resolution: the live product catalogue is loaded from the matching service. Gemini attempts a canonical match constrained to the catalogue; otherwise it identifies a product name.','Stage 7 - Persistence: a Pydantic DemandEventModel is serialised with evidence, confidence, availability and resolution-stage metadata and added to `demand_events`.']: d.add_paragraph(t,style='List Number')
    d.add_paragraph('5. Matching and Catalogue Cache',style='Heading 1')
    d.add_paragraph('`MatchingService` keeps an in-memory canonical catalogue and alias map protected by a threading lock. `refresh_cache()` rebuilds both structures from Firestore. `get_top_candidates()` uses RapidFuzz `WRatio` for the top N candidates. `match_product_stage4()` checks a shop-specific alias first, then scores canonical names and global aliases. The current demand pipeline prefers Gemini’s catalogue-constrained result; the matching methods remain available for compatible/future flows.')
    d.add_paragraph('6. Demand Intelligence and Procurement',style='Heading 1')
    d.add_paragraph('`DemandIntelligenceAgent.process_demand_metrics(shop_id)` is scheduled via FastAPI BackgroundTasks after a successful capture. It retrieves recent shop events, uses pandas to calculate unavailable counts, request frequencies, demand scores and trending products, writes a `demand_summary/{shop_id}` projection, then invokes the procurement workflow. The Procurement Agent reads the summary and the current business-insights document, prompts Gemini for structured recommendations conforming to `ProcurementRecommendation`, and writes `recommendations/{shop_id}`.')
    d.add_paragraph('Recommended production implementation: enqueue a versioned event containing event ID and shop ID; make summary/recommendation writes idempotent; store input/projection versions for traceability; and use bounded retries for model and Firestore failures.')
    d.add_paragraph('7. Business Intelligence Agent',style='Heading 1')
    d.add_paragraph('`BusinessIntelligenceAgent.run_update()` executes at startup and every six hours. It combines configured external context (weather and trend signals where available) with local `data/festivals.json`, then replaces the singleton `business_insights/latest` document. This isolates slow-changing external context from the low-latency demand capture path.')
    d.add_paragraph('8. Firestore Physical Model',style='Heading 1')
    matrix(d,['Collection / path','Essential fields','Writer'],[
        ('products/{id}','canonical_name, aliases, category, brand','Backend catalogue APIs / seed'),('demand_events/{id}','shop_id, product, canonical_product, availability, evidence, confidence, timestamp','Demand Capture Agent'),('demand_summary/{shop_id}','counts, frequencies, scores, trending_products, updated_at','Demand Intelligence Agent'),('business_insights/latest','weather, trends, festivals, updated_at','Business Intelligence Agent'),('recommendations/{shop_id}','recommendations[], updated_at','Procurement Agent'),('procurement_feedback/{id}','shop_id, recommendation_id, feedback, comments, timestamp','Feedback API')],[2400,4700,2260])
    d.add_paragraph('9. Client Integration',style='Heading 1')
    add_bullets(d,['The Flutter client owns capture UI, API calls and presentation; API base configuration is centralised in `frontend/lib/config.dart`.','`api_service.dart` calls the FastAPI backend, while `firestore_service.dart` exposes Firestore-based update streams for dashboard data.','Audio utility files provide web, I/O and stub implementations so the application can build across supported targets.','Locale provider and ARB resources support English/Hindi localisation. Provider manages client-side state; shared preferences supports local settings.'])
    d.add_paragraph('10. Error Handling and Security Details',style='Heading 1')
    matrix(d,['Area','Current behaviour','Production requirement'],[
        ('Validation','Pydantic models validate key API payloads','Define error-code catalogue and contract tests'),('Agent failures','Logs exception; API maps unhandled error to 500','Timeout/retry policy, safe user message, dead-letter handling'),('Credentials','Environment/service-account configuration; mock fallback exists','Secret manager, workload identity, no file-based secrets in deployment'),('Firestore','Admin SDK writes; mock client supports local work','Rules by tenant/role, indexes, audit trails and backups'),('CORS','Prototype permits all origins','Explicit deployed origins and secure credentials policy')],[1700,3660,4000])
    d.add_paragraph('11. Test and Release Strategy',style='Heading 1')
    add_bullets(d,['Maintain unit tests for request models, matching thresholds, score calculation, Firestore mock behaviour and prompt-output parsing.','Add deterministic fixture transcripts covering Hindi/Hinglish variants, unavailable/available/unknown answers, multiple requests and malformed outputs.','Run API contract tests against the mock datastore; run a minimal staging suite against Firestore emulator or isolated project.','Before release, scan for secrets, validate environment configuration, build all intended Flutter targets, and execute one end-to-end capture-to-recommendation smoke test.'])
    d.add_paragraph('12. Known Prototype Gaps and Delivery Plan',style='Heading 1')
    matrix(d,['Gap','Impact','Recommended next action'],[
        ('In-process background tasks/scheduler','Work can be lost on restart and duplicate across replicas','Adopt durable queue and managed scheduler'),('Global CORS wildcard','Not suitable for public deployment','Restrict to known origins'),('Alias confirmation service stub','Endpoint cannot complete intended alias update','Implement per-shop alias persistence with validation'),('No explicit auth/tenancy layer','Data isolation is incomplete','Add authenticated identity and shop-scoped authorization'),('No recommendation evaluation loop','Quality cannot be measured consistently','Join feedback with recommendation versions and analytics')],[3000,2900,3460])
    d.add_paragraph('13. Definition of Done',style='Heading 1'); d.add_paragraph('A production increment is complete when its API contract, persistence changes, Firestore access controls, agent behaviour, tests, observability, and rollback plan are reviewed together. No model-generated recommendation should produce a supplier order without an explicitly authorised shopkeeper action.')
    d.save(ROOT/'LLD_Saudagar_AI.docx')

if __name__ == '__main__':
    hld(); lld()
